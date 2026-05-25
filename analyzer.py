import re
import pdfplumber
from docx import Document

TECH_SKILLS = [
    'python','java','javascript','typescript','react','angular','vue','node','express',
    'django','flask','fastapi','spring','sql','mysql','postgresql','mongodb','redis',
    'docker','kubernetes','aws','azure','gcp','terraform','git','ci/cd','linux',
    'html','css','tailwind','graphql','rest','api','machine learning','deep learning',
    'tensorflow','pytorch','scikit-learn','pandas','numpy','data analysis','nlp',
    'computer vision','spark','hadoop','kafka','elasticsearch','microservices',
    'c++','c#','go','rust','swift','kotlin','php','ruby','scala','r','matlab',
    'tableau','power bi','excel','figma','photoshop','agile','scrum','jira',
    'devops','cloud','serverless','firebase','supabase','netlify','vercel',
    'blockchain','web3','solidity','selenium','pytest','junit','cypress',
    'jenkins','github actions','ansible','nginx','apache','redis','rabbitmq'
]

SOFT_SKILLS = [
    'leadership','communication','teamwork','problem solving','critical thinking',
    'project management','time management','adaptability','creativity','collaboration',
    'analytical','detail-oriented','self-motivated','organized','initiative',
    'presentation','negotiation','mentoring','strategic','planning'
]

ACTION_VERBS = [
    'developed','designed','implemented','managed','led','created','built','optimized',
    'improved','increased','reduced','launched','delivered','coordinated','analyzed',
    'architected','deployed','automated','streamlined','collaborated','mentored',
    'spearheaded','achieved','exceeded','transformed','established'
]

def parse_resume(filepath):
    ext = filepath.rsplit('.', 1)[1].lower()
    text = ''
    try:
        if ext == 'pdf':
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text += t + '\n'
        elif ext == 'docx':
            doc = Document(filepath)
            for para in doc.paragraphs:
                text += para.text + '\n'
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + ' '
                    text += '\n'
    except Exception as e:
        text = f'Error parsing file: {str(e)}'
    return text.strip()

def extract_skills(text):
    text_lower = text.lower()
    found_tech = [s for s in TECH_SKILLS if re.search(r'\b' + re.escape(s) + r'\b', text_lower)]
    found_soft = [s for s in SOFT_SKILLS if re.search(r'\b' + re.escape(s) + r'\b', text_lower)]
    return {'technical': found_tech, 'soft': found_soft}

def extract_contact_info(text):
    email = re.findall(r'[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}', text)
    phone = re.findall(r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]', text)
    linkedin = re.findall(r'linkedin\.com/in/[\w-]+', text, re.IGNORECASE)
    github = re.findall(r'github\.com/[\w-]+', text, re.IGNORECASE)
    return {
        'email': email[0] if email else None,
        'phone': phone[0] if phone else None,
        'linkedin': linkedin[0] if linkedin else None,
        'github': github[0] if github else None
    }

def extract_sections(text):
    sections = {}
    text_lower = text.lower()
    section_patterns = {
        'experience': r'(experience|employment|work history)',
        'education': r'(education|academic|qualification)',
        'skills': r'(skills|expertise|competencies|technologies)',
        'projects': r'(projects|portfolio|works)',
        'certifications': r'(certification|certificate|credential)',
        'summary': r'(summary|objective|profile|about)'
    }
    for section, pattern in section_patterns.items():
        sections[section] = bool(re.search(pattern, text_lower))
    return sections

def count_action_verbs(text):
    text_lower = text.lower()
    found = [v for v in ACTION_VERBS if re.search(r'\b' + v + r'\b', text_lower)]
    return found

def estimate_years_experience(text):
    years = re.findall(r'(\d+)\+?\s*years?\s*(?:of\s*)?(?:experience|exp)', text, re.IGNORECASE)
    if years:
        return max(int(y) for y in years)
    date_ranges = re.findall(r'(20[0-9]{2})\s*[-–]\s*(20[0-9]{2}|present|current)', text, re.IGNORECASE)
    if date_ranges:
        from datetime import datetime
        current_year = datetime.now().year
        total = 0
        for start, end in date_ranges:
            end_year = current_year if end.lower() in ('present', 'current') else int(end)
            total += max(0, end_year - int(start))
        return min(total, 30)
    return 0

def get_improvement_suggestions(text, skills, sections, ats_score):
    suggestions = []
    text_lower = text.lower()

    if not sections.get('summary'):
        suggestions.append({'type': 'warning', 'icon': '📋', 'text': 'Add a professional summary/objective section to make a strong first impression.'})
    if not sections.get('skills'):
        suggestions.append({'type': 'error', 'icon': '⚡', 'text': 'Add a dedicated Skills section — ATS systems scan for this heavily.'})
    if len(text.split()) < 300:
        suggestions.append({'type': 'warning', 'icon': '📝', 'text': 'Resume seems short. Aim for 400–700 words with detailed descriptions.'})
    if not sections.get('certifications'):
        suggestions.append({'type': 'info', 'icon': '🏆', 'text': 'Consider adding certifications to boost credibility and ATS score.'})

    action_verbs = count_action_verbs(text)
    if len(action_verbs) < 5:
        suggestions.append({'type': 'warning', 'icon': '💪', 'text': f'Use more action verbs. Only {len(action_verbs)} found. Try: developed, implemented, led, optimized.'})

    if not re.search(r'\d+%|\$[\d,]+|\d+x|\d+\s*(million|k|thousand)', text_lower):
        suggestions.append({'type': 'warning', 'icon': '📊', 'text': 'Add quantifiable achievements (e.g., "Reduced load time by 40%" or "Led team of 8").'})

    contact = extract_contact_info(text)
    if not contact['linkedin']:
        suggestions.append({'type': 'info', 'icon': '🔗', 'text': 'Add your LinkedIn profile URL to increase professional visibility.'})
    if not contact['github'] and len(skills['technical']) > 3:
        suggestions.append({'type': 'info', 'icon': '💻', 'text': 'Add your GitHub profile to showcase your technical projects.'})

    if len(skills['technical']) < 8:
        suggestions.append({'type': 'error', 'icon': '🛠️', 'text': f'Only {len(skills["technical"])} technical skills detected. Expand your skills section with relevant technologies.'})

    if ats_score >= 80:
        suggestions.append({'type': 'success', 'icon': '🎯', 'text': 'Excellent ATS score! Your resume is well-optimized for applicant tracking systems.'})
    elif ats_score >= 60:
        suggestions.append({'type': 'info', 'icon': '✅', 'text': 'Good ATS score. A few more relevant keywords would push this higher.'})

    return suggestions

def calculate_ats_score(resume_text, job_description):
    score = 0
    text_lower = resume_text.lower()
    sections = extract_sections(resume_text)
    skills = extract_skills(resume_text)
    contact = extract_contact_info(resume_text)
    action_verbs = count_action_verbs(resume_text)

    # Section completeness (25 pts)
    section_scores = {'summary': 5, 'experience': 8, 'education': 5, 'skills': 7}
    for sec, pts in section_scores.items():
        if sections.get(sec):
            score += pts

    # Contact info (10 pts)
    if contact['email']: score += 4
    if contact['phone']: score += 3
    if contact['linkedin']: score += 3

    # Skills (20 pts)
    skill_score = min(20, len(skills['technical']) * 1.5 + len(skills['soft']) * 0.5)
    score += skill_score

    # Action verbs (10 pts)
    score += min(10, len(action_verbs) * 0.8)

    # Quantifiable achievements (10 pts)
    if re.search(r'\d+%|\$[\d,]+|\d+x', text_lower):
        score += 10

    # Length (10 pts)
    word_count = len(resume_text.split())
    if 350 <= word_count <= 800:
        score += 10
    elif 200 <= word_count < 350 or 800 < word_count <= 1000:
        score += 5

    # Job description matching (15 pts)
    if job_description:
        job_lower = job_description.lower()
        job_words = set(re.findall(r'\b[a-z]{3,}\b', job_lower)) - {'the','and','for','are','with','that','this','have','from','they','will','been','were','their','what','your','which','when','there','about','would','could','should','into','also','than','then','some','more','over','after','before'}
        resume_words = set(re.findall(r'\b[a-z]{3,}\b', text_lower))
        if job_words:
            match_ratio = len(job_words & resume_words) / len(job_words)
            score += min(15, match_ratio * 15)

    return round(min(score, 100), 1)

def get_missing_skills(resume_text, job_description):
    if not job_description:
        return []
    job_lower = job_description.lower()
    resume_lower = resume_text.lower()
    job_skills = [s for s in TECH_SKILLS if re.search(r'\b' + re.escape(s) + r'\b', job_lower)]
    missing = [s for s in job_skills if not re.search(r'\b' + re.escape(s) + r'\b', resume_lower)]
    return missing[:15]

def analyze_resume(text):
    skills = extract_skills(text)
    sections = extract_sections(text)
    contact = extract_contact_info(text)
    action_verbs = count_action_verbs(text)
    years_exp = estimate_years_experience(text)
    word_count = len(text.split())
    ats_score = calculate_ats_score(text, '')

    suggestions = get_improvement_suggestions(text, skills, sections, ats_score)

    score_breakdown = {
        'Sections': min(25, sum([5 if sections.get('summary') else 0, 8 if sections.get('experience') else 0, 5 if sections.get('education') else 0, 7 if sections.get('skills') else 0])),
        'Contact Info': (4 if contact['email'] else 0) + (3 if contact['phone'] else 0) + (3 if contact['linkedin'] else 0),
        'Skills': round(min(20, len(skills['technical']) * 1.5 + len(skills['soft']) * 0.5), 1),
        'Action Verbs': round(min(10, len(action_verbs) * 0.8), 1),
        'Achievements': 10 if re.search(r'\d+%|\$[\d,]+|\d+x', text.lower()) else 0,
        'Length': 10 if 350 <= word_count <= 800 else (5 if 200 <= word_count <= 1000 else 0)
    }

    return {
        'skills': skills,
        'sections': sections,
        'contact': contact,
        'action_verbs': action_verbs,
        'years_experience': years_exp,
        'word_count': word_count,
        'suggestions': suggestions,
        'score_breakdown': score_breakdown
    }
